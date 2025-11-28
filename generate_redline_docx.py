from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime

def generate_redlined_docx(clauses, output_path="redlined_output.docx"):
    """Creates a professional DOCX file highlighting clauses by risk level."""
    doc = Document()
    
    # Title
    title = doc.add_heading("AI Contract Review Report", 0)
    title.alignment = 1  # Center alignment
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = 1
    meta.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}\n")
    meta.add_run("AI-Powered Contract Analysis System")
    
    doc.add_paragraph()  # Spacing

    # Summary Section
    summary_heading = doc.add_heading("Executive Summary", 1)
    high_risk = len([c for c in clauses if c.get("risk_level", "").lower() == "high"])
    medium_risk = len([c for c in clauses if c.get("risk_level", "").lower() == "medium"])
    low_risk = len([c for c in clauses if c.get("risk_level", "").lower() == "low"])
    
    summary = doc.add_paragraph()
    summary.add_run("📊 Risk Distribution:\n").bold = True
    summary.add_run(f"• High Risk Clauses: {high_risk}\n")
    summary.add_run(f"• Medium Risk Clauses: {medium_risk}\n") 
    summary.add_run(f"• Low Risk Clauses: {low_risk}\n")
    
    if high_risk > 0:
        warning = doc.add_paragraph()
        warning_run = warning.add_run("⚠️  Attention Required: This contract contains high-risk clauses that need legal review.\n")
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_page_break()

    # Detailed Analysis
    doc.add_heading("Detailed Clause Analysis", 1)
    
    for i, c in enumerate(clauses):
        # Clause Header with Risk Level
        clause_header = doc.add_heading(f"Clause {i+1}", 2)
        
        risk_label = doc.add_paragraph()
        risk_run = risk_label.add_run(f"Risk Level: {c.get('risk_level', 'Unknown')}")
        risk_run.bold = True
        
        # Color code risk level
        risk_color = c.get('highlight_color', 'black')
        if risk_color == "red":
            risk_run.font.color.rgb = RGBColor(220, 38, 38)
        elif risk_color == "orange":
            risk_run.font.color.rgb = RGBColor(245, 158, 11)
        elif risk_color == "green":
            risk_run.font.color.rgb = RGBColor(34, 197, 94)

        # Original Clause
        doc.add_paragraph("Original Clause:").bold = True
        orig_para = doc.add_paragraph(c.get("original_clause", ""))
        orig_para.style = 'Intense Quote'

        # Suggested Clause
        doc.add_paragraph("AI Suggested Improvement:").bold = True
        sug_para = doc.add_paragraph()
        sug_run = sug_para.add_run(c.get("suggested_clause", ""))
        sug_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run("Matching Standard: ").bold = True
        meta_para.add_run(c.get("matched_clause", ""))
        
        if c.get("similarity_score"):
            meta_para.add_run("\nConfidence Score: ").bold = True
            meta_para.add_run(f"{c.get('similarity_score', 0):.2%}")

        # Separator
        if i < len(clauses) - 1:
            doc.add_paragraph("―" * 80)
            doc.add_paragraph()

    # Footer
    doc.add_page_break()
    footer_section = doc.add_heading("Disclaimer", 1)
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("This AI-generated review is for informational purposes only and does not constitute legal advice. Always consult with qualified legal professionals for contract review and final decision-making.")

    try:
        doc.save(output_path)
        print(f"✅ Redlined document saved: {output_path}")
        return output_path
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        return None