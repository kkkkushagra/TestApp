from fastapi import APIRouter, UploadFile, File, HTTPException
from sentence_transformers import SentenceTransformer, util
import torch
import pandas as pd
import docx
import fitz  # PyMuPDF for PDFs
import io

router = APIRouter(prefix="/process_contract", tags=["Process Contract"])

# Load playbook
try:
    df = pd.read_csv("clause_playbook.csv")
except Exception as e:
    raise RuntimeError(f"Failed to load clause playbook: {e}")

# Model
model = SentenceTransformer("all-MiniLM-L6-v2")

# Helper: extract text
def extract_text(file: UploadFile):
    if file.filename.endswith(".pdf"):
        text = ""
        with fitz.open(stream=file.file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text("text")
        return text
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(file.file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload PDF or DOCX.")

# Helper: find best match clause
def analyze_clause(clause):
    query_embedding = model.encode(clause, convert_to_tensor=True)
    clause_embeddings = model.encode(df["standard_clause"].tolist(), convert_to_tensor=True)
    similarity_scores = util.pytorch_cos_sim(query_embedding, clause_embeddings)[0]
    best_idx = int(similarity_scores.argmax())
    best_score = float(similarity_scores[best_idx])

    result = {
        "text": clause.strip(),
        "matched_clause": df["standard_clause"].iloc[best_idx],
        "risk_level": df["Risk_Level"].iloc[best_idx],
        "action_required": df["Action_Required"].iloc[best_idx],
        "similarity_score": round(best_score, 3)
    }
    return result

@router.post("/")
async def process_contract(file: UploadFile = File(...)):
    text = extract_text(file)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Empty document uploaded.")

    # Split clauses by period or newline
    clauses = [c.strip() for c in text.split(".") if len(c.strip()) > 10]

    results = [analyze_clause(clause) for clause in clauses]

    # Create redlined docx
    reviewed_doc = docx.Document()
    for r in results:
        para = reviewed_doc.add_paragraph()
        risk = r["risk_level"]
        run = para.add_run(r["text"])
        if risk == "High":
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        elif risk == "Medium":
            run.font.color.rgb = docx.shared.RGBColor(255, 165, 0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)

        para.add_run(f"  [Risk: {risk} | Action: {r['action_required']}]")

    output_path = "reviewed_contract.docx"
    reviewed_doc.save(output_path)

    return {
        "message": "Contract processed successfully.",
        "total_clauses": len(results),
        "high_risk": sum(r["risk_level"] == "High" for r in results),
        "medium_risk": sum(r["risk_level"] == "Medium" for r in results),
        "low_risk": sum(r["risk_level"] == "Low" for r in results),
        "download_path": output_path
    }
