# redline.py (edited)
from fastapi import APIRouter, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse, FileResponse
import tempfile, os, math, numpy as np, pandas as pd, torch, re, gc, time, shutil, subprocess, logging
from sentence_transformers import SentenceTransformer, util
import fitz, pdfplumber
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import RGBColor
from docx.oxml.text.paragraph import CT_P
from typing import List
import html as _html
import uuid

router = APIRouter(prefix="/redline", tags=["Redlining"])
logger = logging.getLogger("redline")
logging.basicConfig(level=logging.INFO)

def sanitize_for_json(data):
    if isinstance(data, float):
        if math.isnan(data) or math.isinf(data):
            return 0.0
        return data
    elif isinstance(data, dict):
        return {k: sanitize_for_json(v) for k,v in data.items()}
    elif isinstance(data, list):
        return [sanitize_for_json(v) for v in data]
    elif isinstance(data, np.ndarray):
        return sanitize_for_json(data.tolist())
    else:
        return data

# Load playbook
PLAYBOOK_PATH = "clause_playbook.csv"
if not os.path.exists(PLAYBOOK_PATH):
    raise RuntimeError(f"Clause playbook not found at {PLAYBOOK_PATH}")

playbook = pd.read_csv(PLAYBOOK_PATH)
for col in ["standard_clause","Risk_Level","Action_Required"]:
    if col not in playbook.columns:
        playbook[col] = ""
playbook["standard_clause"] = playbook["standard_clause"].fillna("").astype(str)
playbook["Risk_Level"] = playbook["Risk_Level"].fillna("Low").astype(str)
playbook["Action_Required"] = playbook["Action_Required"].fillna("").astype(str)
playbook_clauses = playbook["standard_clause"].tolist()
logger.info(f"✅ Loaded playbook ({len(playbook_clauses)} entries)")

# Load sentence-transformers model (LegalBERT)
logger.info("🔄 Loading sentence-transformers model...")
model = SentenceTransformer("nlpaueb/legal-bert-base-uncased")
playbook_emb = model.encode(playbook_clauses, convert_to_tensor=True)
# convert possible NaNs
try:
    playbook_emb = torch.nan_to_num(playbook_emb, nan=0.0)
except Exception:
    pass
logger.info("✅ Playbook embeddings ready.")

# Extraction functions
def extract_text_from_pdf(path: str, ocr_max_pages: int = 3) -> str:
    frags = []
    try:
        with fitz.open(path) as doc:
            for pnum, page in enumerate(doc, start=1):
                try:
                    t = page.get_text("text") or ""
                    if t.strip():
                        frags.append(f"\n--- Page {pnum} ---\n" + t)
                except Exception as e:
                    logger.exception("fitz page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("PyMuPDF failed: %s", e)
    try:
        with pdfplumber.open(path) as pdf:
            for pnum, page in enumerate(pdf.pages, start=1):
                try:
                    t = page.extract_text() or ""
                    if t.strip():
                        frags.append(f"\n--- Page {pnum} (pdfplumber) ---\n" + t)
                except Exception as e:
                    logger.exception("pdfplumber page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("pdfplumber failed: %s", e)
    try:
        pages = convert_from_path(path, dpi=200, first_page=1, last_page=ocr_max_pages)
        for pnum, pil in enumerate(pages, start=1):
            try:
                t = pytesseract.image_to_string(pil, config="--oem 1 --psm 6")
                if t.strip():
                    frags.append(f"\n--- Page {pnum} (OCR) ---\n" + t)
            except Exception as e:
                logger.exception("OCR page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("OCR failed: %s", e)
    return ""

def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception as e:
        logger.exception("DOCX extract failed:")
        return ""

def split_into_clauses(text: str) -> List[str]:
    if not text or not text.strip(): return []
    text = text.replace("\r", "\n")
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'-\n', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[\.\?\!])\s+', text)
    clauses, current = [], ""
    for s in sentences:
        s = s.strip()
        if not s: continue
        if len(s.split()) < 6 and current:
            current += " " + s
        else:
            if current: clauses.append(current.strip())
            current = s
    if current: clauses.append(current.strip())
    clauses = [c for c in clauses if len(c.split()) >= 3]
    if not clauses: clauses = [text]
    return clauses

# Create tracked-changes-like docx: original strikethrough + suggested insert
def create_redlined_doc(results, original_filename, output_path):
    try:
        doc = Document()
        doc.add_heading("AI Contract Review Report", 0).alignment = 1
        doc.add_paragraph(f"Generated on {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Source file: {original_filename}")
        doc.add_page_break()

        doc.add_heading("Summary", level=1)
        total = len(results)
        high = sum(1 for r in results if r["risk_level"].lower() == "high")
        medium = sum(1 for r in results if r["risk_level"].lower() == "medium")
        low = sum(1 for r in results if r["risk_level"].lower() == "low")
        doc.add_paragraph(f"Total clauses analyzed: {total}")
        doc.add_paragraph(f"High risk: {high} — Medium risk: {medium} — Low risk: {low}")
        doc.add_page_break()

        doc.add_heading("Detailed Clause Analysis", level=1)
        for i, r in enumerate(results):
            doc.add_heading(f"Clause {i+1}", level=2)
            rp = doc.add_paragraph()
            rp.add_run("Risk Level: ").bold = True
            run = rp.add_run(r.get("risk_level","Unknown"))
            rl = r.get("risk_level","").lower()
            if rl == "high": run.font.color.rgb = RGBColor(255,0,0)
            elif rl == "medium": run.font.color.rgb = RGBColor(245,158,11)
            else: run.font.color.rgb = RGBColor(34,197,94)

            doc.add_paragraph("Original Clause (struck-through):").bold = True
            p = doc.add_paragraph()
            orig_run = p.add_run(r.get("original_clause",""))
            try:
                orig_run.font.strike = True
            except Exception:
                pass

            doc.add_paragraph("Suggested Clause (inserted):").bold = True
            p2 = doc.add_paragraph()
            sug_run = p2.add_run(r.get("suggested_clause",""))
            try:
                sug_run.font.color.rgb = RGBColor(0, 112, 192)  # blue inserted text
                try:
                    sug_run.underline = True
                except Exception:
                    pass
            except Exception:
                pass

            doc.add_paragraph("Explanation:").bold = True
            doc.add_paragraph(r.get("explanation",""))
            doc.add_paragraph("―" * 80)

        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.exception("create_redlined_doc failed:")
        return None

# -------------------------
#  Replacement helpers
# -------------------------
def _replace_text_preserve_runs(paragraph, search_text, replace_text):
    """
    Replace first occurrence of search_text in paragraph.text with replace_text,
    attempting to preserve formatting by reconstructing runs using the first run's style.
    Returns True if a replacement was made.
    """
    if not search_text:
        return False
    full = paragraph.text
    idx = full.find(search_text)
    if idx == -1:
        return False

    # Gather run formatting to try to reuse
    first_run = paragraph.runs[0] if paragraph.runs else None
    # Build new text
    new_text = full[:idx] + replace_text + full[idx + len(search_text):]

    # Clear existing runs
    for r in paragraph.runs:
        try:
            # r.clear() is not always present; set text to empty
            r.text = ""
        except Exception:
            try:
                # attempt to remove run via xml if necessary
                r._r.getparent().remove(r._r)
            except Exception:
                pass

    # Add a new run with new_text
    new_run = paragraph.add_run(new_text)
    try:
        if first_run:
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            try:
                new_run.font.size = first_run.font.size
            except Exception:
                pass
            try:
                new_run.font.name = first_run.font.name
            except Exception:
                pass
            try:
                new_run.font.underline = first_run.font.underline
            except Exception:
                pass
            try:
                new_run.font.color.rgb = first_run.font.color.rgb
            except Exception:
                pass
    except Exception:
        pass

    return True

def _replace_in_tables(doc, replacements):
    changed = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for orig, repl in replacements.items():
                        if orig and orig in para.text:
                            _replace_text_preserve_runs(para, orig, repl)
                            changed = True
    return changed

def _replace_in_paragraphs(doc, replacements):
    changed = False
    for para in doc.paragraphs:
        for orig, repl in replacements.items():
            if orig and orig in para.text:
                _replace_text_preserve_runs(para, orig, repl)
                changed = True
    return changed

def _replace_in_headers_footers(doc, replacements):
    changed = False
    try:
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for orig, repl in replacements.items():
                    if orig and orig in para.text:
                        _replace_text_preserve_runs(para, orig, repl)
                        changed = True
            footer = section.footer
            for para in footer.paragraphs:
                for orig, repl in replacements.items():
                    if orig and orig in para.text:
                        _replace_text_preserve_runs(para, orig, repl)
                        changed = True
    except Exception:
        pass
    return changed

# Helper: convert PDF -> DOCX using LibreOffice (soffice). Returns converted docx path or None.
def convert_pdf_to_docx_with_soffice(pdf_path: str, out_dir: str = None, timeout: int = 60) -> str | None:
    try:
        if not out_dir:
            out_dir = os.path.dirname(pdf_path) or tempfile.gettempdir()
        os.makedirs(out_dir, exist_ok=True)

        # locate soffice
        soffice = os.environ.get("SOFFICE_PATH") or shutil.which("soffice")
        if not soffice:
            logger.warning("soffice not found in PATH. Set SOFFICE_PATH or install LibreOffice.")
            return None

        logger.info("Running soffice conversion: %s --headless --convert-to docx --outdir %s %s", soffice, out_dir, pdf_path)
        # run conversion
        proc = subprocess.run([soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, pdf_path],
                              stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
        if proc.returncode != 0:
            logger.warning("soffice conversion exit %s; stderr=%s stdout=%s", proc.returncode, proc.stderr.decode("utf-8", errors="ignore"), proc.stdout.decode("utf-8", errors="ignore"))
            return None

        # Determine converted filename
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        # LibreOffice typically writes base + .docx in outdir
        candidate = os.path.join(out_dir, base + ".docx")
        if os.path.exists(candidate):
            logger.info("soffice conversion succeeded -> %s", candidate)
            return os.path.abspath(candidate)

        # try to find any docx file in out_dir created recently
        now = time.time()
        recent = []
        for fn in os.listdir(out_dir):
            if fn.lower().endswith(".docx"):
                full = os.path.join(out_dir, fn)
                # if created in the last 120 seconds assume it's our file
                if now - os.path.getmtime(full) < 120:
                    recent.append(full)
        if recent:
            logger.info("soffice conversion found docx file -> %s", recent[0])
            return os.path.abspath(recent[0])

        logger.warning("soffice conversion completed but could not find .docx output in %s", out_dir)
        return None
    except subprocess.TimeoutExpired:
        logger.exception("soffice conversion timed out")
        return None
    except Exception:
        logger.exception("soffice conversion failed")
        return None

# Helper: wrap fragments into HTML (first occurrence safe match)
def wrap_fragments_in_html(html_text: str, fragments: List[dict]) -> str:
    out = html_text
    for fr in fragments:
        orig = (fr.get("original") or "").strip()
        if not orig:
            continue
        seed = orig[:200]
        try:
            pattern = re.escape(seed)
            m = re.search(pattern, out)
            if not m:
                pattern2 = re.escape(seed[:80])
                m = re.search(pattern2, out, flags=re.IGNORECASE)
            if m:
                matched_text = out[m.start():m.end()]
                span = f'<span class="risky" data-clause-id="{fr.get("clause_id")}" data-risk="{fr.get("risk_level", "")}">{_html.escape(matched_text)}</span>'
                out = out[:m.start()] + span + out[m.end():]
            else:
                continue
        except Exception:
            continue
    return out

# Helper: build basic HTML from page blocks (re-uses extraction functions)
def build_simple_html_from_text_blocks(page_blocks: List[str]) -> str:
    body = "<div class='document'>\n"
    for i, p in enumerate(page_blocks):
        body += f"<div class='page' data-page='{i+1}'>\n"
        paras = [para.strip() for para in re.split(r'\n{2,}', p) if para.strip()]
        for para in paras:
            body += f"<p class='pblock'>{_html.escape(para)}</p>\n"
        body += "</div>\n"
    body += "</div>"
    full = (
        "<html><head><meta charset='utf-8'/>"
        "<style>"
        ".risky{ text-decoration: underline wavy red; text-underline-offset: 3px; cursor:pointer; } "
        ".risky[data-risk='Medium']{ text-decoration-color: orange; } "
        ".page{ padding:12px 18px; margin-bottom:20px; border-bottom:1px solid rgba(255,255,255,0.02); } "
        ".pblock{ margin:6px 0; white-space:pre-wrap; font-family: system-ui, -apple-system, 'Segoe UI', Roboto, 'Helvetica Neue', Arial; font-size:14px; color:#e6e6e6 }"
        "</style></head><body>" + body + "</body></html>"
    )
    return full

@router.post("/convert/")
async def convert_to_html(file: UploadFile = File(None), payload: dict = Body(None)):
    """
    Convert uploaded PDF/DOCX (or server-side path) into simple HTML with risky fragments wrapped as <span class='risky'>.
    Accepts either:
      - form-data file upload (file)
      - JSON body { "path": "/absolute/server/path/to/file.pdf" }
    Returns:
      { html, clauses, summary, original_docx }
    """
    source_path = None
    tmp_path = None
    try:
        if file:
            ext = os.path.splitext(file.filename or "upload")[1].lower() or ".pdf"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmp_path = tmp.name
            try:
                while True:
                    chunk = await file.read(1024 * 1024)
                    if not chunk:
                        break
                    tmp.write(chunk)
            finally:
                try: tmp.flush(); tmp.close()
                except: pass
            source_path = tmp_path
        elif payload and isinstance(payload, dict) and payload.get("path"):
            p = payload.get("path")
            if not os.path.exists(p):
                raise HTTPException(status_code=400, detail=f"Path not found: {p}")
            source_path = p
        else:
            raise HTTPException(status_code=400, detail="Provide either form-data file or JSON { 'path': '/server/path/to/file' }")

        _, ext = os.path.splitext(source_path)
        ext = ext.lower()

        page_blocks = []
        if ext == ".pdf":
            try:
                with fitz.open(source_path) as doc:
                    for pnum, page in enumerate(doc, start=1):
                        try:
                            blocks = page.get_text("blocks")
                            blocks.sort(key=lambda b: (round(b[1]), round(b[0])))
                            page_text_parts = []
                            for b in blocks:
                                txt = (b[4] or "").strip()
                                if txt:
                                    page_text_parts.append(txt)
                            page_blocks.append("\n\n".join(page_text_parts))
                        except Exception:
                            try:
                                t = page.get_text("text") or ""
                                page_blocks.append(t)
                            except Exception:
                                page_blocks.append("")
            except Exception as e:
                txt = extract_text_from_pdf(source_path, ocr_max_pages=3)
                if txt:
                    page_blocks = re.split(r'\n--- Page \d+ ---\n', txt)
                    page_blocks = [p for p in page_blocks if p and p.strip()]
                else:
                    page_blocks = [txt]
        elif ext == ".docx":
            txt = extract_text_from_docx(source_path)
            page_blocks = [txt]
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type for convert (use PDF or DOCX)")

        full_text = "\n\n".join([p for p in page_blocks if p and p.strip()])
        clauses_raw = split_into_clauses(full_text)

        clause_objs = []
        for idx, ctext in enumerate(clauses_raw):
            if not ctext or len(ctext.strip()) < 6:
                continue
            try:
                emb = model.encode(ctext, convert_to_tensor=True)
                emb = torch.nan_to_num(emb, nan=0.0)
                sims = util.pytorch_cos_sim(emb, playbook_emb)[0]
                sims = torch.nan_to_num(sims, nan=0.0)
                best_idx = int(torch.argmax(sims))
                score = float(sims[best_idx])
            except Exception:
                best_idx = 0; score = 0.0

            matched_standard = playbook_clauses[best_idx]
            matched_risk = str(playbook.loc[best_idx, "Risk_Level"]) if "Risk_Level" in playbook.columns else "Low"
            matched_action = str(playbook.loc[best_idx, "Action_Required"]) if "Action_Required" in playbook.columns else ""

            rl = matched_risk.lower()
            if rl == "high":
                suggested = matched_standard or ctext
                explanation = f"High-risk clause. Action required: {matched_action or 'Legal Review'}"
                highlight_color = "red"
            elif rl == "medium":
                suggested = matched_standard or ctext
                explanation = f"Medium-risk clause. Action recommended: {matched_action or 'Review'}"
                highlight_color = "orange"
            else:
                suggested = ctext
                explanation = "Low-risk clause. No change recommended."
                highlight_color = "green"

            clause_objs.append({
                "clause_id": f"c{idx+1}",
                "original_clause": ctext,
                "suggested_clause": suggested,
                "matched_clause": matched_standard,
                "risk_level": str(matched_risk),
                "action_required": matched_action,
                "similarity_score": score,
                "highlight_color": highlight_color,
                "explanation": explanation
            })

        fragments_to_wrap = [
            {"clause_id": o["clause_id"], "original": o["original_clause"], "risk_level": o["risk_level"]}
            for o in clause_objs if o["risk_level"].lower() in ("high", "medium")
        ]

        html_body = build_simple_html_from_text_blocks(page_blocks)
        html_with_spans = wrap_fragments_in_html(html_body, fragments_to_wrap)

        words = re.findall(r'\w+', full_text)
        total_words = len(words)
        uniq_clauses = len(set([c.strip() for c in clauses_raw]))
        dup_count = max(0, len(clauses_raw) - uniq_clauses)
        risk_score = sum(3 if r["risk_level"].lower() == "high" else (2 if r["risk_level"].lower() == "medium" else 1) for r in clause_objs)

        summary = {
            "total_clauses": len(clause_objs),
            "high_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "high"),
            "medium_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "medium"),
            "low_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "low"),
            "total_words": total_words,
            "unique_clauses": uniq_clauses,
            "duplicate_clauses": dup_count,
            "risk_score": risk_score
        }

        original_docx_token = None
        try:
            saved_uploaded_dir = "uploaded_contracts"
            os.makedirs(saved_uploaded_dir, exist_ok=True)
            unique_token = f"{uuid.uuid4().hex}{ext}"
            saved_path = os.path.join(saved_uploaded_dir, unique_token)
            try:
                shutil.copy(source_path, saved_path)
                # If the saved file is a PDF, attempt to convert to docx and store that converted docx path as original_docx
                if ext == ".pdf":
                    # Try convert to a temp outdir
                    conv_outdir = os.path.join(tempfile.gettempdir(), f"soffice_out_{uuid.uuid4().hex}")
                    os.makedirs(conv_outdir, exist_ok=True)
                    converted = convert_pdf_to_docx_with_soffice(saved_path, out_dir=conv_outdir)
                    if converted and os.path.exists(converted):
                        # Copy converted docx to uploaded_contracts and return that absolute path
                        conv_dest = os.path.join(saved_uploaded_dir, f"{uuid.uuid4().hex}.docx")
                        shutil.copy(converted, conv_dest)
                        original_docx_token = os.path.abspath(conv_dest)
                        logger.info("Persisted converted docx for formatting-preservation: %s", original_docx_token)
                    else:
                        # If conversion failed, keep saved_path (pdf) as persisted source for best-effort fallback
                        original_docx_token = os.path.abspath(saved_path)
                        logger.warning("Could not convert uploaded PDF to DOCX for formatting-preservation; saved original PDF at %s", original_docx_token)
                else:
                    original_docx_token = os.path.abspath(saved_path)
            except Exception as e:
                logger.exception("Warning: could not persist source file:")
                original_docx_token = None
        except Exception as e:
            logger.exception("Warning while trying to persist uploaded original:")
            original_docx_token = None

        return JSONResponse(content=sanitize_for_json({
            "html": html_with_spans,
            "clauses": clause_objs,
            "summary": summary,
            "original_docx": original_docx_token
        }))

    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

@router.post("/generate/")
async def generate_from_clauses(payload: dict = Body(...)):
    """
    Generate a DOCX that contains full document text reconstructed from 'clauses',
    but with risky clauses replaced by their suggested_clause (if provided).
    Expected payload:
      {
        "filename": "optional_name",
        "clauses": [ { clause_id, original_clause, suggested_clause, risk_level, explanation }, ... ],
        "original_docx": "/absolute/path/to/original.docx"   # OPTIONAL, server-side path returned by /convert/
      }
    Returns:
      { download_link: "/redline/download/yourfile.docx" }
    """
    fname = payload.get("filename") or f"edited_{int(time.time())}"
    safe_name = re.sub(r'[^\w\-_\.]', '_', str(fname))
    clauses = payload.get("clauses") or []
    original_docx = payload.get("original_docx") or None

    if not isinstance(clauses, list) or len(clauses) == 0:
        raise HTTPException(status_code=400, detail="No clauses provided to generate doc.")

    replacements = {}
    for c in clauses:
        risk = (c.get("risk_level") or "Low").lower()
        if risk in ("high", "medium"):
            orig = (c.get("original_clause") or "").strip()
            sug = (c.get("suggested_clause") or orig).strip()
            if orig:
                replacements[orig] = sug

    out_dir = "generated_reports"; os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{safe_name}.docx")

    try:
        # If original_docx provided and exists on server, attempt in-place structured replacement
        if original_docx and isinstance(original_docx, str) and os.path.exists(original_docx):
            try:
                # If provided path is a PDF file, try to convert it to docx first
                _, orig_ext = os.path.splitext(original_docx)
                if orig_ext.lower() == ".pdf":
                    logger.info("original_docx points to a PDF; attempting to convert with soffice before structured replace.")
                    conv_outdir = os.path.join(tempfile.gettempdir(), f"soffice_out_{uuid.uuid4().hex}")
                    os.makedirs(conv_outdir, exist_ok=True)
                    converted = convert_pdf_to_docx_with_soffice(original_docx, out_dir=conv_outdir)
                    if converted and os.path.exists(converted):
                        # use converted docx
                        logger.info("Converted PDF -> DOCX for generation: %s", converted)
                        original_docx = converted
                    else:
                        logger.warning("Conversion of original PDF to DOCX failed; will attempt fallback generation.")

                # Only attempt structured replace if file is a .docx
                if original_docx.lower().endswith(".docx") and os.path.exists(original_docx):
                    doc = Document(original_docx)
                    made_change = False

                    p_changed = _replace_in_paragraphs(doc, replacements)
                    t_changed = _replace_in_tables(doc, replacements)
                    hf_changed = _replace_in_headers_footers(doc, replacements)
                    made_change = p_changed or t_changed or hf_changed

                    if made_change:
                        doc.save(out_path)
                    else:
                        # Nothing matched: fallback to assembling text
                        assembled_lines = []
                        for c in clauses:
                            orig = c.get("original_clause", "") or ""
                            risk = (c.get("risk_level") or "Low").lower()
                            suggested = c.get("suggested_clause", "") or orig
                            if risk in ("high", "medium"):
                                assembled_lines.append(suggested.strip())
                            else:
                                assembled_lines.append(orig.strip())
                        full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])

                        new_doc = Document()
                        new_doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                        new_doc.add_paragraph("")  # spacer
                        for para in full_text.split("\n\n"):
                            p = new_doc.add_paragraph()
                            p.add_run(para)
                        new_doc.save(out_path)
                else:
                    # Not a docx or conversion failed: fallback assembly
                    assembled_lines = []
                    for c in clauses:
                        orig = c.get("original_clause", "") or ""
                        risk = (c.get("risk_level") or "Low").lower()
                        suggested = c.get("suggested_clause", "") or orig
                        if risk in ("high", "medium"):
                            assembled_lines.append(suggested.strip())
                        else:
                            assembled_lines.append(orig.strip())
                    full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])

                    doc = Document()
                    doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                    doc.add_paragraph("")  # spacer
                    for para in full_text.split("\n\n"):
                        p = doc.add_paragraph()
                        p.add_run(para)
                    doc.save(out_path)

            except Exception as e:
                logger.exception("Error while trying structured replace in original_docx:")
                # Fallback: create plain docx from clauses
                assembled_lines = []
                for c in clauses:
                    orig = c.get("original_clause", "") or ""
                    risk = (c.get("risk_level") or "Low").lower()
                    suggested = c.get("suggested_clause", "") or orig
                    if risk in ("high", "medium"):
                        assembled_lines.append(suggested.strip())
                    else:
                        assembled_lines.append(orig.strip())
                full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])

                new_doc = Document()
                new_doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                new_doc.add_paragraph("")  # spacer
                for para in full_text.split("\n\n"):
                    p = new_doc.add_paragraph()
                    p.add_run(para)
                new_doc.save(out_path)
        else:
            # No original_docx available — create simple DOCX by concatenation
            assembled_lines = []
            for c in clauses:
                orig = c.get("original_clause") or ""
                risk = (c.get("risk_level") or "Low").lower()
                suggested = c.get("suggested_clause") or orig
                if risk in ("high", "medium"):
                    assembled_lines.append(suggested.strip())
                else:
                    assembled_lines.append(orig.strip())
            full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])

            doc = Document()
            doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")  # spacer

            for para in full_text.split("\n\n"):
                p = doc.add_paragraph()
                p.add_run(para)

            doc.save(out_path)
    except Exception as e:
        logger.exception("generate_from_clauses -> failed to build DOCX:")
        raise HTTPException(status_code=500, detail="Failed to create edited DOCX.")

    return JSONResponse(content=sanitize_for_json({
        "download_link": f"/redline/download/{os.path.basename(out_path)}"
    }))

@router.get("/download/{filename}")
def download_redlined(filename: str):
    path = os.path.join("generated_reports", filename)
    if not os.path.exists(path):
        raise HTTPException(404, "File not found")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

# ---------- END: ADDITIONAL ENDPOINTS ----------
